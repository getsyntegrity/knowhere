import io
import os
from pathlib import Path

import pandas as pd
from docx import Document
from lxml import etree

TEST_RUNTIME_DIR = Path(__file__).resolve().parent / ".tmp_docx_toc_alignment"
TEST_RUNTIME_DIR.mkdir(exist_ok=True)
(TEST_RUNTIME_DIR / "font.ttf").touch(exist_ok=True)
(TEST_RUNTIME_DIR / "chromedriver").touch(exist_ok=True)

os.environ.setdefault("DS_KEY", "test")
os.environ.setdefault("DS_URL", "https://example.com")
os.environ.setdefault("S3_BUCKET_NAME", "test-bucket")
os.environ.setdefault("S3_ACCESS_KEY_ID", "test-key")
os.environ.setdefault("S3_SECRET_ACCESS_KEY", "test-secret")
os.environ.setdefault("S3_TEMP_PATH", str(TEST_RUNTIME_DIR))
os.environ.setdefault("USERS_DATA_PATH", str(TEST_RUNTIME_DIR))
os.environ.setdefault("DATABASE_URL", f"sqlite:///{TEST_RUNTIME_DIR / 'test.db'}")
os.environ.setdefault("SECRET_KEY", "test-secret-key")
os.environ.setdefault("TMP_PATH", str(TEST_RUNTIME_DIR))
os.environ.setdefault("FONT_PATH", str(TEST_RUNTIME_DIR / "font.ttf"))
os.environ.setdefault("CHROMEDRIVER_PATH", str(TEST_RUNTIME_DIR / "chromedriver"))

from app.services.document_parser import doc_parser
from app.services.document_parser import layout_parser
from app.services.document_parser.layout_parser import format_toc_context_for_llm
from app.services.document_parser.toc_parser import (
    build_docx_toc_hierarchies,
    detect_doc_tocs,
)


DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def _make_paragraph(text: str):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(buffer)
    return Document(io.BytesIO(buffer.getvalue())).paragraphs[0]


def test_detect_doc_tocs_extracts_numeric_level_from_toc_style() -> None:
    elem = etree.fromstring(
        """
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:pPr>
            <w:pStyle w:val="TOC2" />
          </w:pPr>
          <w:r>
            <w:t>1.1 适用范围</w:t>
          </w:r>
        </w:p>
        """
    )

    toc_info = detect_doc_tocs(elem, DOCX_NS)

    assert toc_info["is_style"] is True
    assert toc_info["toc_level"] == 2
    assert toc_info["style_name"] == "TOC2"


def test_detect_doc_tocs_uses_outline_level_for_unnumbered_toc_style() -> None:
    elem = etree.fromstring(
        """
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:pPr>
            <w:pStyle w:val="TOC" />
            <w:outlineLvl w:val="1" />
          </w:pPr>
          <w:r>
            <w:t>第一章 总则</w:t>
          </w:r>
        </w:p>
        """
    )

    toc_info = detect_doc_tocs(elem, DOCX_NS)

    assert toc_info["is_style"] is True
    assert toc_info["toc_level"] == 2
    assert toc_info["outline_level"] == 2
    assert toc_info["style_name"] == "TOC"


def test_build_docx_toc_hierarchies_preserves_levels_and_groups_areas() -> None:
    toc_hierarchies = build_docx_toc_hierarchies(
        [
            (
                1,
                "目录",
                "TOC-AREA",
                {"toc_level": None, "toc_style_name": "TOCHeading"},
            ),
            (2, "第一章 总则", "TOC-AREA", {"toc_level": 1, "toc_style_name": "TOC1"}),
            (3, "1.1 适用范围", "TOC-AREA", {"toc_level": 2, "toc_style_name": "TOC2"}),
            (4, "正文开始", "PTXT", None),
            (5, "附录A", "TOC-AREA", {"toc_level": 1, "toc_style_name": "TOC1"}),
        ]
    )

    assert len(toc_hierarchies) == 2
    assert toc_hierarchies[0]["toc_range"] == (1, 3)
    assert "目录" not in toc_hierarchies[0]["toc_with_level"]
    assert "第一章 总则" in toc_hierarchies[0]["toc_with_level"]
    assert toc_hierarchies[0]["toc_tree"] == {"第一章 总则": {"1.1 适用范围": {}}}
    assert toc_hierarchies[1]["toc_tree"] == {"附录A": {}}


def test_build_docx_toc_hierarchies_falls_back_to_indent_for_plain_toc_styles() -> None:
    toc_hierarchies = build_docx_toc_hierarchies(
        [
            (
                1,
                "目录",
                "TOC-AREA",
                {"toc_level": None, "toc_style_name": "TOC", "toc_left_indent": 0},
            ),
            (
                2,
                "总则",
                "TOC-AREA",
                {"toc_level": None, "toc_style_name": "TOC", "toc_left_indent": 0},
            ),
            (
                3,
                "适用范围",
                "TOC-AREA",
                {"toc_level": None, "toc_style_name": "TOC", "toc_left_indent": 240},
            ),
            (4, "正文开始", "PTXT", None),
        ]
    )

    assert len(toc_hierarchies) == 1
    assert "目录" not in toc_hierarchies[0]["toc_with_level"]
    assert "总则" in toc_hierarchies[0]["toc_with_level"]
    assert "适用范围" in toc_hierarchies[0]["toc_with_level"]
    assert toc_hierarchies[0]["toc_tree"] == {"总则": {"适用范围": {}}}


def test_format_toc_context_for_llm_accepts_markdown_table_payload() -> None:
    toc_context = [
        {
            "toc_range": (2, 3),
            "toc_with_level": (
                "| id | heading | level | reason |\n"
                "| --- | --- | --- | --- |\n"
                "| 2 | 第一章 总则 | 1 | POS [] NEG [] |"
            ),
        }
    ]

    formatted = format_toc_context_for_llm(toc_context)

    assert "TOC 1 (source rows 2-3):" in formatted
    assert "| id | heading | level | reason |" in formatted
    assert "第一章 总则" in formatted


def test_parse_docx_passes_structured_toc_context_to_pred_titles(
    monkeypatch,
    tmp_path: Path,
) -> None:
    paragraph = _make_paragraph("第一章 总则")
    block_tuples = [
        (1, "目录", "TOC-AREA", {"toc_level": None, "toc_style_name": "TOCHeading"}),
        (2, "第一章 总则", "TOC-AREA", {"toc_level": 1, "toc_style_name": "TOC1"}),
        (3, paragraph, "PTXT", None),
        (4, "正文内容", "PTXT", None),
    ]
    captured = {}

    monkeypatch.setattr(
        doc_parser, "load_file_bytes", lambda *args, **kwargs: b"docx-bytes"
    )
    monkeypatch.setattr(doc_parser, "iter_block_items", lambda doc_data: block_tuples)

    def fake_pred_titles(heading_infos, **kwargs):
        captured["heading_infos"] = heading_infos
        captured["kwargs"] = kwargs
        return pd.DataFrame(
            [
                {
                    "id": 3,
                    "heading": "第一章 总则",
                    "level": 1,
                    "reason": "POS [] NEG []",
                }
            ]
        )

    monkeypatch.setattr(doc_parser, "pred_titles", fake_pred_titles)

    parsed_structure, _ = doc_parser.parse_docx(
        "sample.docx",
        llm_paras={
            "smart_title_parse": True,
            "doc_type": "manual",
            "summary_image": False,
            "summary_table": False,
            "model_name": "deepseek-chat",
            "hierarchy_model_name": "qwen3.6-flash",
        },
        output_dir=str(tmp_path),
        filename="sample.docx",
    )

    assert parsed_structure["content"][0]["heading"] == "第一章 总则"
    assert captured["kwargs"]["first_toc_ele_num"] == 1
    assert captured["kwargs"]["toc_hierarchies"][0]["toc_range"] == (1, 2)
    assert "第一章 总则" in captured["kwargs"]["toc_hierarchies"][0]["toc_with_level"]
    assert captured["kwargs"]["model_name"] == "qwen3.6-flash"
    assert (tmp_path / "toc_hierarchies.json").exists()


def test_pred_titles_uses_zone_parsing_for_docx_and_prefers_earliest_toc_boundary(
    monkeypatch,
) -> None:
    calls = []

    def fake_filter_doc_headings(*args, **kwargs):
        return pd.DataFrame(
            [
                {"id": 0, "heading": "封面", "level": -1, "reason": "cover"},
                {"id": 4, "heading": "第一章 总则", "level": 1, "reason": "l1"},
                {"id": 7, "heading": "1.1 适用范围", "level": 2, "reason": "l2"},
                {"id": 10, "heading": "第二章 术语", "level": 1, "reason": "l1"},
            ]
        )

    def fake_est_hierarchies_naive(raw_preds, proceed_smart=True, output_dir=None):
        return raw_preds.copy()

    def fake_est_hierarchies_llm(
        raw_preds, prompt_limt, toc_hierarchies=None, **kwargs
    ):
        calls.append(
            {
                "ids": raw_preds["id"].tolist(),
                "toc_range": toc_hierarchies[0]["toc_range"],
            }
        )
        result = raw_preds.copy()
        result["level"] = [1] * len(result)
        return result

    monkeypatch.setattr(layout_parser, "filter_doc_headings", fake_filter_doc_headings)
    monkeypatch.setattr(
        layout_parser, "est_hierarchies_naive", fake_est_hierarchies_naive
    )
    monkeypatch.setattr(layout_parser, "est_hierarchies_llm", fake_est_hierarchies_llm)

    heading_preds = layout_parser.pred_titles(
        infos=[],
        doc_type="docx",
        toc_hierarchies=[
            {
                "toc_range": (2, 3),
                "toc_with_level": "toc-1",
                "toc_tree": {"第一章 总则": {}},
            },
            {
                "toc_range": (8, 9),
                "toc_with_level": "toc-2",
                "toc_tree": {"第二章 术语": {}},
            },
        ],
        smart_parse=True,
        first_toc_ele_num=1,
    )

    assert calls == [
        {"ids": [4, 7], "toc_range": (2, 3)},
        {"ids": [10], "toc_range": (8, 9)},
    ]
    cover_row = heading_preds.loc[heading_preds["id"] == 0].iloc[0]
    assert cover_row["level"] == -1
