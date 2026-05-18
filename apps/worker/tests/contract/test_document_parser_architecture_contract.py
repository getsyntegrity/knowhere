from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pandas as pd


def test_parse_input_builds_typed_llm_parameters(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.orchestration.parse_input import (
        ParseInput,
        ParseOptions,
    )
    from app.services.document_parser.orchestration.parse_session import (
        build_parse_session,
    )

    monkeypatch.setattr(
        "app.services.document_parser.orchestration.parse_session.profile_document",
        lambda *_args, **_kwargs: SimpleNamespace(
            file_type="pdf",
            page_count=3,
            atlas_candidate=False,
            doc_category="generic",
            summary=lambda: "profile",
            reasoning="test",
        ),
    )

    parse_input = ParseInput(
        file_full_path=str(tmp_path / "sample.pdf"),
        filename="sample.pdf",
        output_dir=str(tmp_path),
        internal_output_filename="internal.pdf",
        job_id="job-1",
        kb_dir="Default_Root",
        options=ParseOptions(
            doc_type="auto",
            llm_histories=7,
            smart_title_parse=False,
            summary_image=False,
            summary_table=True,
            summary_txt=False,
            stopwords=["the"],
            add_frag_desc="fragment",
        ),
        s3_key="uploads/sample.pdf",
    )

    session = build_parse_session(parse_input)

    assert session.base_llm_paras == {
        "llm_histories": 7,
        "smart_title_parse": False,
        "summary_image": False,
        "summary_table": True,
        "summary_txt": False,
        "stopwords": ["the"],
        "doc_type": "auto",
        "frag_desc": "fragment",
        "model_name": session.base_llm_paras["model_name"],
        "hierarchy_model_name": session.base_llm_paras["hierarchy_model_name"],
    }
    assert session.relative_root == "Default_Root/sample.pdf"


def test_document_format_router_uses_adapters(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.orchestration.format_router import (
        DocumentFormat,
        get_document_parse_adapter,
        resolve_document_format,
    )
    from app.services.document_parser.orchestration.parse_input import ParseInput
    from app.services.document_parser.orchestration.parse_output import ParseOutput
    from app.services.document_parser.orchestration.parse_session import ParseSession
    from app.services.document_parser.orchestration.route_parse import route_document_parse

    assert resolve_document_format("/tmp/report.PDF") == DocumentFormat.PDF
    assert resolve_document_format("/tmp/report.docx") == DocumentFormat.DOCX
    assert get_document_parse_adapter(DocumentFormat.PDF).document_format == DocumentFormat.PDF

    parsed_df = pd.DataFrame([{"content": "ok"}])

    monkeypatch.setattr(
        "app.services.document_parser.pdf_parser.parse_pdfs",
        lambda *_args, **_kwargs: parsed_df,
    )

    profile = SimpleNamespace(route="standard", doc_category="generic")
    parse_input = ParseInput(
        file_full_path=str(tmp_path / "report.pdf"),
        filename="report.pdf",
        output_dir=str(tmp_path),
        internal_output_filename="report.pdf",
    )
    session = ParseSession.from_input(
        parse_input=parse_input,
        base_llm_paras={},
        full_output_dir=str(tmp_path),
        profile=profile,
        relative_root="Default_Root/report.pdf",
    )

    output = route_document_parse(session)

    assert output == ParseOutput(output_dir=str(tmp_path), parsed_df=parsed_df)
    assert output.as_legacy_tuple() == (str(tmp_path), parsed_df)


def test_parse_pipeline_owns_route_and_postprocess_contract(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.orchestration.parse_input import ParseInput
    from app.services.document_parser.orchestration.parse_output import ParseOutput
    from app.services.document_parser.orchestration.parse_pipeline import (
        ParsePipelineResult,
        run_parse_pipeline,
    )

    routed_df = pd.DataFrame([{"content": "routed"}])
    processed_df = pd.DataFrame([{"content": "processed"}])
    routed_output_dir = str(tmp_path / "routed")

    monkeypatch.setattr(
        "app.services.document_parser.orchestration.parse_session.profile_document",
        lambda *_args, **_kwargs: SimpleNamespace(
            file_type="pdf",
            page_count=1,
            atlas_candidate=False,
            doc_category="generic",
            summary=lambda: "profile",
            reasoning="test",
        ),
    )
    monkeypatch.setattr(
        "app.services.document_parser.orchestration.parse_pipeline.route_document_parse",
        lambda _session: ParseOutput(output_dir=routed_output_dir, parsed_df=routed_df),
    )

    def fake_postprocess(output_dir: str, parsed_df: pd.DataFrame | None) -> pd.DataFrame:
        assert output_dir == routed_output_dir
        assert parsed_df is routed_df
        return processed_df

    monkeypatch.setattr(
        "app.services.document_parser.orchestration.parse_pipeline.apply_parse_postprocess",
        fake_postprocess,
    )

    result = run_parse_pipeline(
        ParseInput(
            file_full_path=str(tmp_path / "contract.pdf"),
            filename="contract.pdf",
            output_dir=str(tmp_path),
            internal_output_filename="contract.pdf",
        )
    )

    assert result == ParsePipelineResult(
        output_dir=routed_output_dir,
        parsed_df=processed_df,
    )
    assert result.as_legacy_tuple() == (routed_output_dir, processed_df)


def test_parser_service_keeps_legacy_tuple_compatibility(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    import app.services.document_parser.parse_service as parse_service
    from app.services.document_parser.orchestration.parse_output import ParseOutput

    parsed_df = pd.DataFrame([{"content": "legacy"}])
    expected_output = ParseOutput(output_dir=str(tmp_path), parsed_df=parsed_df)

    monkeypatch.setattr(
        parse_service,
        "run_parse_pipeline",
        lambda _parse_input: expected_output,
    )

    output = parse_service.checkerboard_parse_output(
        file_full_path=str(tmp_path / "legacy.pdf"),
        filename="legacy.pdf",
        output_dir=str(tmp_path),
        internal_output_filename="legacy.pdf",
    )
    legacy_tuple = parse_service.checkerboard_inject_parse(
        file_full_path=str(tmp_path / "legacy.pdf"),
        filename="legacy.pdf",
        output_dir=str(tmp_path),
        internal_output_filename="legacy.pdf",
    )

    assert output is expected_output
    assert legacy_tuple == (str(tmp_path), parsed_df)


def test_rendered_pdf_transform_centralizes_temporary_pdf_cleanup(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.rendered_pdf_transform import (
        parse_rendered_pdf_bytes,
    )

    parsed_df = pd.DataFrame([{"content": "pptx"}])
    seen_pdf_bytes: list[bytes] = []

    def fake_parse_pdfs(pdf_path: str, **_kwargs: Any) -> pd.DataFrame:
        seen_pdf_bytes.append(Path(pdf_path).read_bytes())
        assert Path(pdf_path).exists()
        return parsed_df

    monkeypatch.setattr(
        "app.services.document_parser.rendered_pdf_transform.parse_pdfs",
        fake_parse_pdfs,
    )
    monkeypatch.setattr(
        "app.services.document_parser.rendered_pdf_transform.render_pdf_to_image_pdf",
        lambda pdf_bytes: pdf_bytes,
    )

    actual_df = parse_rendered_pdf_bytes(
        pdf_bytes=b"rendered",
        filename="slides.pptx",
        output_dir=str(tmp_path),
        base_llm_paras={},
        relative_root="Default_Root/slides.pptx",
        rendered_pdf_s3_key="transforms/job.pdf",
    )

    assert actual_df is parsed_df
    assert seen_pdf_bytes == [b"rendered"]
    assert not (tmp_path / "_pptx_tmp.pdf").exists()


def test_heading_hierarchy_module_wraps_prediction(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.heading_hierarchy import (
        HeadingHierarchyInput,
        predict_heading_hierarchy,
    )

    expected_df = pd.DataFrame(
        [{"id": 1, "heading": "Intro", "level": 1, "reason": "test"}]
    )
    captured: dict[str, Any] = {}

    def fake_pred_titles(*args: Any, **kwargs: Any) -> pd.DataFrame:
        captured["args"] = args
        captured["kwargs"] = kwargs
        return expected_df

    monkeypatch.setattr(
        "app.services.document_parser.heading_hierarchy.pred_titles",
        fake_pred_titles,
    )

    actual_df = predict_heading_hierarchy(
        HeadingHierarchyInput(
            infos=[(1, "Intro")],
            doc_type="md",
            smart_parse=True,
            model_name="hierarchy-model",
            output_dir=str(tmp_path),
            layout_json_path=str(tmp_path / "layout.json"),
        )
    )

    assert actual_df is expected_df
    assert captured["kwargs"]["doc_type"] == "md"
    assert captured["kwargs"]["smart_parse"] is True
    assert captured["kwargs"]["model_name"] == "hierarchy-model"


def test_parser_row_builder_owns_dataframe_column_order(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.parser_rows import ParsedRow, ParsedRowsBuilder

    builder = ParsedRowsBuilder()
    builder.append(
        ParsedRow(
            content="chunk text",
            path="Default_Root/doc/Section",
            type="text",
            keywords="alpha;beta",
            summary="summary",
            know_id="chunk-1",
            tokens="alpha->beta",
            connectto="",
            page_nums="1,2",
            addtime="now",
        )
    )

    parsed_df = builder.to_dataframe()

    assert list(parsed_df.columns) == [
        "content",
        "path",
        "type",
        "length",
        "keywords",
        "summary",
        "know_id",
        "tokens",
        "connectto",
        "addtime",
        "page_nums",
    ]
    assert parsed_df.iloc[0].to_dict() == {
        "content": "chunk text",
        "path": "Default_Root/doc/Section",
        "type": "text",
        "length": len("chunk text"),
        "keywords": "alpha;beta",
        "summary": "summary",
        "know_id": "chunk-1",
        "tokens": "alpha->beta",
        "connectto": "",
        "addtime": "now",
        "page_nums": "1,2",
    }


def test_inline_asset_module_builds_image_and_table_rows(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.inline_asset import (
        build_image_asset_row,
        build_table_asset_row,
    )

    image_row = build_image_asset_row(
        content="\nImage summary\n[images/image-1.png]\n",
        relative_path="images/image-1.png",
        summary="image-1\nImage summary",
        know_id="image-1",
        addtime="now",
        page_nums="3",
    )
    table_row = build_table_asset_row(
        content="<table></table>",
        relative_path="tables/table-1.html",
        summary="table-1\nTable summary",
        keywords="column",
        know_id="table-1",
        addtime="now",
        page_nums="4",
    )

    assert image_row.type == "image"
    assert image_row.path == "images/image-1.png"
    assert image_row.summary == "image-1\nImage summary"
    assert table_row.type == "table"
    assert table_row.path == "tables/table-1.html"
    assert table_row.keywords == "column"


def test_table_asset_writer_creates_table_row_and_html_file(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.table_asset_writer import (
        TableAssetInput,
        write_table_asset,
    )

    row = write_table_asset(
        TableAssetInput(
            html="<table><tr><td>A</td></tr></table>",
            output_dir=str(tmp_path),
            table_name="table-1",
            summary="table-1",
            keywords="A",
            know_id="table-1",
            addtime="now",
        )
    )

    assert (tmp_path / "tables" / "table-1.html").read_text(encoding="utf-8")
    assert row.type == "table"
    assert row.path == "tables/table-1.html"
    assert row.content == "<table><tr><td>A</td></tr></table>"


def test_docx_asset_store_owns_asset_filesystem_lifecycle(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.docx_asset_store import DocxAssetStore

    store = DocxAssetStore(str(tmp_path))
    (tmp_path / "images").mkdir()
    (tmp_path / "tables").mkdir()
    (tmp_path / "images" / "stale.png").write_bytes(b"stale")
    (tmp_path / "tables" / "stale.html").write_text("stale", encoding="utf-8")

    store.reset()
    image_asset = store.write_image("image-1 raw", ".png", b"image")
    renamed_asset = store.rename_image(image_asset, "image-1 final")
    table_asset = store.write_table("table-1 final", "<table></table>")

    assert not (tmp_path / "images" / "stale.png").exists()
    assert not (tmp_path / "tables" / "stale.html").exists()
    assert renamed_asset.relative_path == "images/image-1 final.png"
    assert (tmp_path / "images" / "image-1 final.png").read_bytes() == b"image"
    assert table_asset.relative_path == "tables/table-1 final.html"
    assert (tmp_path / "tables" / "table-1 final.html").read_text(
        encoding="utf-8"
    ) == "<table></table>"


def test_docx_block_stream_emits_document_ordered_blocks(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.docx_block_stream import iter_block_items
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Intro")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    document.save(docx_path)

    block_events = list(iter_block_items(docx_path.read_bytes()))

    assert block_events[0][0] == 1
    assert isinstance(block_events[0][1], Paragraph)
    assert block_events[0][1].text == "Intro"
    assert block_events[0][2] == "PTXT"
    assert isinstance(block_events[1][1], Table)
    assert block_events[1][2] == "TABLE"


def test_html_table_modules_separate_docx_and_dataframe_rendering(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.dataframe_html_renderer import df2html
    from app.services.document_parser.docx_table_html import table2html
    from docx import Document

    docx_path = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.save(docx_path)

    loaded_table = Document(str(docx_path)).tables[0]
    docx_html = table2html(loaded_table, cell_image_map={(0, 1): "image summary"})

    dataframe_html = df2html(
        pd.DataFrame([["North", "North", 3]], columns=["Region", "Group", "Value"]),
        row_header_cols=2,
    )

    assert docx_html == (
        "<table border='1'><tr><td>A</td>"
        "<td>B<br/><em>image summary</em></td></tr></table>"
    )
    assert '<th scope="row" colspan="2">North</th>' in dataframe_html
    assert "<td>3</td>" in dataframe_html


def test_table_text_parser_owns_markdown_table_text_contract(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.table_text_parser import (
        df2md,
        extract_tables_by_forms,
        identify_tables,
        sanitize_table_name_from_header,
    )

    markdown_table = "\n".join(
        [
            "| Product | Revenue |",
            "| --- | ---: |",
            "| Notebook | 42 |",
        ]
    )

    is_table, table_form, _tables = identify_tables("| Product | Revenue |")
    table_html = extract_tables_by_forms(markdown_table, form="md")
    markdown_output = df2md(pd.DataFrame([{"City": "北京", "Value": 7}]))

    assert is_table is True
    assert table_form == "md"
    assert table_html is not None
    assert "<th>Product</th>" in table_html
    assert "<td>Notebook</td>" in table_html
    assert "---" not in table_html
    assert sanitize_table_name_from_header("A | Revenue | Revenue | 市场") == (
        "Revenue 市场"
    )
    assert "| City | Value |" in markdown_output
    assert "| 北京 | 7     |" in markdown_output


def test_table_frame_parser_owns_dataframe_table_contract(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.table_frame_parser import (
        parse_tb_contents,
        parse_tb_keywords,
        postprocess_tb,
    )

    raw_frame = pd.DataFrame(
        [["North\nAmerica", 42, None]],
        columns=["Region\nName", "Revenue", None],
    )

    normalized_frame = postprocess_tb(raw_frame, drop=True)
    paths, table_html = parse_tb_contents(
        normalized_frame,
        parent_dic={"budget.xlsx": {"Visible": {}}},
        file_name="budget.xlsx",
        sheet_name="Visible",
    )
    keywords = parse_tb_keywords(normalized_frame)

    assert normalized_frame.columns.tolist() == ["RegionName", "Revenue"]
    assert "NorthAmerica" in table_html
    assert "RegionName" in keywords
    assert "Revenue" in keywords
    assert "budget.xlsx/Visible/RegionName" in paths
    assert "budget.xlsx/Visible/Revenue" in paths


def test_markdown_table_asset_module_owns_table_asset_contract(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.markdown_table_asset import (
        MarkdownTableAssetRequest,
        build_markdown_table_asset,
    )
    from app.services.document_parser.markdown_deferred_task import (
        TableDeferredSummaryTask,
    )

    table_dir = tmp_path / "tables"
    table_dir.mkdir()
    table_html = (
        "<table><thead><tr><th>Product</th><th>Revenue</th></tr></thead>"
        "<tbody><tr><td>Notebook</td><td>42</td></tr></tbody></table>"
    )

    asset = build_markdown_table_asset(
        MarkdownTableAssetRequest(
            table_html=table_html,
            table_dir=str(table_dir),
            table_count=3,
            timestamp="now",
            current_page_number=9,
            summary_table=True,
            row_index=7,
        )
    )

    assert asset.content_item == f"\n[{asset.relative_path}]\n"
    assert asset.row_values[1] == asset.relative_path
    assert asset.row_values[2] == "table"
    assert asset.row_values[5] == "table-3"
    assert asset.row_values[10] == "9"
    assert asset.deferred_task == TableDeferredSummaryTask(
        row_index=7,
        table_html=table_html,
        table_dir=str(table_dir),
        table_name=Path(asset.relative_path).stem,
        table_count=2,
    )
    assert "border='1'" in (tmp_path / asset.relative_path).read_text(
        encoding="utf-8"
    )


def test_markdown_image_asset_module_owns_image_materialization_contract(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    from app.services.document_parser.markdown_image_asset import (
        MarkdownImageAssetRequest,
        build_markdown_image_name,
        build_markdown_image_asset,
    )
    from app.services.document_parser.markdown_deferred_task import (
        ImageDeferredSummaryTask,
    )

    image_dir = tmp_path / "images"
    image_dir.mkdir()
    source_image = tmp_path / "raw.png"
    source_image.write_bytes(b"same pixels")
    seen_images: dict[str, dict[str, str]] = {}

    image_asset = build_markdown_image_asset(
        MarkdownImageAssetRequest(
            output_dir=str(tmp_path),
            image_dir=str(image_dir),
            image_path=str(source_image),
            image_name=build_markdown_image_name(
                image_count=2,
                last_context="Revenue Chart",
            ),
            image_count=2,
            last_context="Revenue Chart",
            image_summary="Sales by region",
            timestamp="now",
            current_page_number=8,
            seen_images=seen_images,
            summary_image=True,
            row_index=4,
        )
    )

    assert image_asset.content_item is not None
    assert image_asset.row_values is not None
    assert image_asset.cache_key is not None
    assert image_asset.cache_entry is not None
    assert image_asset.should_advance_image_count is True
    assert image_asset.row_values[1] == "images/image-2-Revenue Ch.png"
    assert image_asset.row_values[2] == "image"
    assert image_asset.row_values[5] == "image-2\nSales by region"
    assert image_asset.row_values[10] == "8"
    assert image_asset.deferred_task == ImageDeferredSummaryTask(
        row_index=4,
        relative_path="images/image-2-Revenue Ch.png",
        image_dir=str(image_dir),
        image_name="image-2-Revenue Ch",
        image_suffix=".png",
    )
    assert (tmp_path / "images" / "image-2-Revenue Ch.png").read_bytes() == (
        b"same pixels"
    )
    assert not source_image.exists()

    seen_images[image_asset.cache_key] = image_asset.cache_entry
    duplicate_source = tmp_path / "duplicate.png"
    duplicate_source.write_bytes(b"same pixels")

    duplicate_asset = build_markdown_image_asset(
        MarkdownImageAssetRequest(
            output_dir=str(tmp_path),
            image_dir=str(image_dir),
            image_path=str(duplicate_source),
            image_name=build_markdown_image_name(
                image_count=3,
                last_context="Other Chart",
            ),
            image_count=3,
            last_context="Other Chart",
            image_summary=None,
            timestamp="now",
            current_page_number=9,
            seen_images=seen_images,
            summary_image=True,
            row_index=5,
        )
    )

    assert duplicate_asset.content_item == image_asset.content_item
    assert duplicate_asset.row_values is not None
    assert duplicate_asset.row_values[1] == "images/image-2-Revenue Ch.png"
    assert duplicate_asset.deferred_task is None
    assert duplicate_asset.should_advance_image_count is False
    assert not duplicate_source.exists()


def test_mineru_modules_separate_client_and_task_polling(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.mineru_client import get_mineru_headers
    from app.services.document_parser.mineru_task_polling import (
        get_batch_status,
        get_polling_interval_for_state,
    )

    assert get_mineru_headers("token") == {
        "Content-Type": "application/json",
        "Authorization": "Bearer token",
    }
    assert get_batch_status({"data": {"extract_result": [{"state": "done"}]}}) == {
        "state": "done"
    }
    assert get_batch_status({"data": {"extract_result": {"state": "failed"}}}) == {
        "state": "failed"
    }
    assert get_polling_interval_for_state("pending", 2) == 8.0
    assert get_polling_interval_for_state("running", 2) == 10.0
    assert get_polling_interval_for_state("waiting-file", 2) == 15.0


def test_doc_profile_model_owns_profile_contract_and_metadata(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    import json

    from app.services.document_parser.doc_profile_model import (
        DocProfile,
        save_profile_metadata,
    )

    profile = DocProfile(
        file_type="pdf",
        route="fast",
        decision_band="safe_fast",
        page_count=3,
        avg_text_density=123.4,
        avg_image_coverage=0.05,
        page_details=[{"page": 1}],
        sample_text="hidden",
    )

    save_profile_metadata(profile, str(tmp_path))
    saved_profile = json.loads((tmp_path / "profile.json").read_text(encoding="utf-8"))

    assert "page_details" not in saved_profile
    assert "sample_text" not in saved_profile
    assert saved_profile["file_type"] == "pdf"
    assert "route=fast" in profile.summary()


def test_doc_profiler_dispatches_pdf_to_pdf_profile_module(
    worker_contract_environment: None,
    monkeypatch: Any,
) -> None:
    from app.services.document_parser.doc_profile_model import DocProfile
    from app.services.document_parser.doc_profiler import profile_document

    called_paths: list[str] = []

    def fake_profile_pdf(path: str) -> DocProfile:
        called_paths.append(path)
        return DocProfile(file_type="pdf", page_count=2)

    monkeypatch.setattr(
        "app.services.document_parser.doc_profiler.profile_pdf",
        fake_profile_pdf,
    )

    pdf_profile = profile_document("/tmp/input.bin", filename="report.pdf")
    docx_profile = profile_document("/tmp/input.bin", filename="report.docx")

    assert called_paths == ["/tmp/input.bin"]
    assert pdf_profile.file_type == "pdf"
    assert docx_profile.file_type == "docx"
    assert docx_profile.route == "standard"


def test_excel_structure_parser_is_table_structure_seam(
    worker_contract_environment: None,
    tmp_path: Path,
) -> None:
    import openpyxl

    from app.services.document_parser.excel_structure_parser import (
        parse_excel_structure,
    )

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Budget"
    worksheet["A1"] = "Region"
    worksheet["B1"] = "Value"
    worksheet["A2"] = "North"
    worksheet["B2"] = 10
    workbook_path = tmp_path / "budget.xlsx"
    workbook.save(workbook_path)

    parsed_sheets = parse_excel_structure(str(workbook_path), split_subtables=False)

    assert list(parsed_sheets.keys()) == ["Budget"]
    assert parsed_sheets["Budget"].attrs["row_header_cols"] >= 0
    assert "North" in parsed_sheets["Budget"].astype(str).to_string()


def test_heading_hierarchy_exposes_candidate_and_tree_modules(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.heading_candidates import filter_markdown_headings
    from app.services.document_parser.heading_tree import cleanup_heading_tree

    candidates = filter_markdown_headings(["# Intro", "body", "## Detail"])
    cleaned = cleanup_heading_tree(
        pd.DataFrame(
            [
                {"id": 0, "heading": "Intro", "level": 1, "reason": ""},
                {"id": 2, "heading": "Detail", "level": 2, "reason": ""},
            ]
        )
    )

    assert candidates[["id", "heading", "level"]].to_dict("records")[0] == {
        "id": 0,
        "heading": "Intro",
        "level": 1,
    }
    assert cleaned["heading"].tolist() == ["Intro", "Detail"]


def test_heading_llm_executor_owns_prompt_execution_and_fallback(
    worker_contract_environment: None,
    monkeypatch: Any,
) -> None:
    from app.services.document_parser.heading_llm_executor import (
        execute_llm_heading_hierarchy,
    )

    monkeypatch.setenv("KB_LAYOUT_LLM_COMPACT_INPUT", "true")
    raw_preds = pd.DataFrame(
        [
            {"id": 0, "heading": "body", "level": -1, "reason": ""},
            {"id": 1, "heading": "Intro", "level": -2, "reason": "POS [1] NEG [0]"},
            {"id": 2, "heading": "body", "level": -1, "reason": ""},
        ]
    )
    judged_prompts: list[pd.DataFrame] = []
    saved_files: list[str] = []

    def fake_hierarchy_judge(
        df: pd.DataFrame,
        *_args: Any,
        **_kwargs: Any,
    ) -> list[dict[str, int]]:
        judged_prompts.append(df.copy())
        return [{"id": 1, "level": 1}]

    def unexpected_fallback(_df: pd.DataFrame) -> pd.DataFrame:
        raise AssertionError("fallback should not run")

    actual_df = execute_llm_heading_hierarchy(
        raw_preds=raw_preds,
        prompt_limt=4000,
        hierarchy_judge=fake_hierarchy_judge,
        fallback_hierarchy=unexpected_fallback,
        save_intermediate_csv=lambda _df, _output_dir, filename: saved_files.append(
            filename
        ),
        model_name="hierarchy-model",
    )

    assert actual_df["level"].tolist() == [-1, 1, -1]
    assert "Intro" in judged_prompts[0]["heading"].tolist()
    assert judged_prompts[0]["heading"].tolist().count("[1 BODY LINES]") == 2
    assert saved_files == ["preds_3_llm_base", "preds_4_llm_final"]

    body_only = pd.DataFrame(
        [{"id": 0, "heading": "body", "level": -1, "reason": ""}]
    )
    skipped_df = execute_llm_heading_hierarchy(
        raw_preds=body_only,
        prompt_limt=4000,
        hierarchy_judge=lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("LLM should not run without heading candidates")
        ),
        fallback_hierarchy=unexpected_fallback,
        save_intermediate_csv=lambda *_args: None,
    )

    assert skipped_df["level"].tolist() == [-1]


def test_markdown_deferred_summary_module_updates_rows_and_refs(
    worker_contract_environment: None,
    monkeypatch: Any,
    tmp_path: Path,
) -> None:
    import app.services.document_parser.markdown_deferred_summary as deferred_summary
    from app.services.document_parser.markdown_deferred_summary import (
        MarkdownDeferredSummaryInput,
        apply_markdown_deferred_summaries,
    )
    from app.services.document_parser.markdown_deferred_task import (
        ImageDeferredSummaryTask,
        TableDeferredSummaryTask,
        TextDeferredSummaryTask,
    )

    image_dir = tmp_path / "images"
    table_dir = tmp_path / "tables"
    image_dir.mkdir()
    table_dir.mkdir()
    (image_dir / "image-3-old.png").write_bytes(b"image")
    (table_dir / "table-0 old.html").write_text("<table></table>", encoding="utf-8")

    rows: list[list[str | int]] = [
        [
            "[images/image-3-old.png]",
            "images/image-3-old.png",
            "image",
            24,
            "",
            "image-3",
            "image-id",
            "",
            "",
            "now",
            "",
        ],
        [
            "[tables/table-0 old.html]",
            "tables/table-0 old.html",
            "table",
            25,
            "",
            "table-0",
            "table-id",
            "",
            "",
            "now",
            "",
        ],
        [
            "long text",
            "Root/Text",
            "text",
            9,
            "",
            "",
            "text-id",
            "",
            "",
            "now",
            "",
        ],
    ]

    monkeypatch.setattr(deferred_summary, "_get_vision_client", lambda: object())
    monkeypatch.setattr(
        deferred_summary,
        "ask_image",
        lambda *_args, **_kwargs: "Better Image\nImage summary",
    )

    def fake_extract(text: str, **_kwargs: Any) -> tuple[str, str, str]:
        if "<table" in text:
            return "Better Table", "table-keyword", "table summary"
        return "Text", "text-keyword", "text summary"

    monkeypatch.setattr(
        deferred_summary, "extract_title_keywords_summary", fake_extract
    )

    apply_markdown_deferred_summaries(
        MarkdownDeferredSummaryInput(
            rows=rows,
            tasks=[
                ImageDeferredSummaryTask(
                    row_index=0,
                    relative_path="images/image-3-old.png",
                    image_dir=str(image_dir),
                    image_name="image-3-old",
                    image_suffix=".png",
                ),
                TableDeferredSummaryTask(
                    row_index=1,
                    table_html="<table></table>",
                    table_dir=str(table_dir),
                    table_name="table-0 old",
                    table_count=0,
                ),
                TextDeferredSummaryTask(row_index=2, content="long text"),
            ],
            output_dir=str(tmp_path),
        )
    )

    assert rows[0][1] == "images/image-3-Better Image.png"
    assert "[images/image-3-Better Image.png]" in str(rows[0][0])
    assert rows[0][5] == "image-3\nImage summary"
    assert (image_dir / "image-3-Better Image.png").exists()
    assert rows[1][1] == "tables/table-0 Better Table.html"
    assert rows[1][4] == "table-keyword"
    assert rows[1][5] == "table-0\ntable summary"
    assert (table_dir / "table-0 Better Table.html").exists()
    assert rows[2][4] == "text-keyword"
    assert rows[2][5] == "text summary"


def test_toc_modules_separate_docx_detection_from_hierarchy_payloads(
    worker_contract_environment: None,
) -> None:
    from app.services.document_parser.toc_docx import infer_toc_level_from_text
    from app.services.document_parser.toc_hierarchy import build_toc_hierarchy_payload

    assert infer_toc_level_from_text("1.2 Scope") == 2
    payload = build_toc_hierarchy_payload(
        [
            {"id": 3, "heading": "1 Overview", "level": 1},
            {"id": 4, "heading": "1.1 Detail", "level": 2},
        ],
        toc_range=(3, 4),
        scan_range=(3, 5),
    )

    assert payload is not None
    assert payload["toc_range"] == (3, 4)
    assert payload["scan_range"] == (3, 5)
    assert payload["toc_tree"] == {"1 Overview": {"1.1 Detail": {}}}


def test_format_adapters_do_not_expose_lazy_any_wrappers(
    worker_contract_environment: None,
) -> None:
    import app.services.document_parser.orchestration.format_adapters as format_adapters

    wrapper_names = [
        "parse_fragment",
        "parse_texts",
        "parse_md",
        "parse_image",
        "parse_pdfs",
        "parse_docx",
        "convert_doc2dics",
        "doc_to_docx",
        "xls_to_xlsx",
        "parse_xlsx",
        "parse_pptx",
    ]

    assert not any(hasattr(format_adapters, wrapper_name) for wrapper_name in wrapper_names)
