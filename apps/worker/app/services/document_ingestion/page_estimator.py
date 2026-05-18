"""
Page estimator for worker-side Document Ingestion billing.

Calculates page counts for billing based on:
- PDF: Physical page count from metadata
- PPTX: Slide count
- Text-based (DOC, DOCX, TXT, MD, JSON): Word-based estimation using count_cn_en
- Spreadsheet-based (XLS, XLSX): Row-based estimation
"""

import math
import tempfile
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path

from shared.core.logging import logger
from shared.utils.text_utils import count_cn_en

WORDS_PER_PAGE = 500
ROWS_PER_PAGE = 50


@dataclass(frozen=True)
class WorkloadEstimate:
    page_count: int
    method: str
    fallback_reason: str | None = None

    @property
    def used_fallback(self) -> bool:
        return self.fallback_reason is not None


class PageEstimator:
    """Estimate page count for billing purposes."""

    @classmethod
    def estimate(cls, file_path: str) -> int:
        """Estimate the billable page count for a file."""
        return cls.estimate_workload(file_path).page_count

    @classmethod
    def estimate_workload(cls, file_path: str) -> WorkloadEstimate:
        """Estimate billable workload while keeping fallback policy explicit."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return cls._estimate_with_fallback("pdf_metadata", cls._count_pdf, file_path)
        if suffix == ".pptx":
            return cls._estimate_with_fallback("pptx_slides", cls._count_pptx, file_path)
        if suffix == ".doc":
            return cls._estimate_with_fallback("doc_conversion", cls._count_doc, file_path)
        if suffix == ".docx":
            return cls._estimate_with_fallback("docx_words", cls._count_docx, file_path)
        if suffix == ".xls":
            return cls._estimate_with_fallback("xls_conversion", cls._count_xls, file_path)
        if suffix == ".xlsx":
            return cls._estimate_with_fallback("xlsx_rows", cls._count_xlsx, file_path)
        if suffix in [".txt", ".md", ".json", ".fragment"]:
            return cls._estimate_with_fallback("text_words", cls._count_text, file_path)
        if suffix in [".png", ".jpg", ".jpeg"]:
            return WorkloadEstimate(page_count=1, method="image_default")

        fallback_reason = f"unknown_file_type:{suffix or '<none>'}"
        logger.warning(
            f"Unknown file type for billing: {suffix}, defaulting to 1 page"
        )
        return WorkloadEstimate(
            page_count=1,
            method="unknown_file_type",
            fallback_reason=fallback_reason,
        )

    @classmethod
    def _estimate_with_fallback(
        cls,
        method: str,
        estimator: Callable[[str], int],
        file_path: str,
    ) -> WorkloadEstimate:
        try:
            page_count = max(1, estimator(file_path))
            return WorkloadEstimate(page_count=page_count, method=method)
        except ImportError as exc:
            fallback_reason = f"missing_dependency:{exc.name or exc}"
            logger.warning(
                f"Missing dependency for {method} page estimation, defaulting to 1 page: {exc}"
            )
        except Exception as exc:
            fallback_reason = f"{method}_error:{type(exc).__name__}"
            logger.error(f"Error estimating pages for {file_path}: {exc}")

        return WorkloadEstimate(
            page_count=1,
            method=method,
            fallback_reason=fallback_reason,
        )

    @classmethod
    def _count_pdf(cls, file_path: str) -> int:
        """Estimate pages for PDF using physical page count."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        return len(reader.pages)

    @classmethod
    def _count_pptx(cls, file_path: str) -> int:
        """Estimate pages for PPTX using slide count."""
        from pptx import Presentation

        presentation = Presentation(file_path)
        return len(presentation.slides)

    @classmethod
    def _count_docx(cls, file_path: str) -> int:
        """Estimate pages for DOCX using word-based counting."""
        from docx import Document

        document = Document(file_path)
        total_text = ""

        for paragraph in document.paragraphs:
            total_text += paragraph.text + " "

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_text += cell.text + " "

        word_count = count_cn_en(total_text)
        return math.ceil(word_count / WORDS_PER_PAGE)

    @classmethod
    def _count_doc(cls, file_path: str) -> int:
        """Estimate pages for DOC by converting it to DOCX first."""
        from app.services.document_parser.conversion.legacy_converter import doc_to_docx

        with tempfile.TemporaryDirectory(prefix="page-estimator-doc-") as temp_dir:
            converted_path, _ = doc_to_docx(file_path, temp_dir)
            return cls._count_docx(converted_path)

    @classmethod
    def _count_xlsx(cls, file_path: str) -> int:
        """Estimate pages for XLSX using row count."""
        import pandas as pd

        workbook = pd.ExcelFile(file_path)
        total_rows = 0

        for sheet_name in workbook.sheet_names:
            dataframe = pd.read_excel(workbook, sheet_name=sheet_name)
            total_rows += len(dataframe)

        return math.ceil(total_rows / ROWS_PER_PAGE)

    @classmethod
    def _count_xls(cls, file_path: str) -> int:
        """Estimate pages for XLS by converting it to XLSX first."""
        from app.services.document_parser.conversion.legacy_converter import xls_to_xlsx

        with tempfile.TemporaryDirectory(prefix="page-estimator-xls-") as temp_dir:
            converted_path, _ = xls_to_xlsx(file_path, temp_dir)
            return cls._count_xlsx(converted_path)

    @classmethod
    def _count_text(cls, file_path: str) -> int:
        """Estimate pages for text-like files using word-based counting."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            content = file.read()

        word_count = count_cn_en(content)
        return math.ceil(word_count / WORDS_PER_PAGE)
