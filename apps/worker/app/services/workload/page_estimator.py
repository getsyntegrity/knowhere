"""
Page Estimator Service

Calculates page counts for billing based on:
- PDF: Physical page count from metadata
- PPTX: Slide count
- Text-based (DOC, DOCX, TXT, MD, JSON): Word-based estimation using count_cn_en
- Spreadsheet-based (XLS, XLSX): Row-based estimation

Supported file types:
- .pdf, .doc, .docx, .pptx, .xls, .xlsx
- .txt, .md, .json, .fragment
- .png, .jpg, .jpeg
"""

import math
import tempfile
from pathlib import Path

from shared.core.logging import logger

# Constants for page estimation
WORDS_PER_PAGE = 500  # Chinese chars + English words + numbers per page
ROWS_PER_PAGE = 50  # For Excel files


from shared.utils.text_utils import count_cn_en


class PageEstimator:
    """
    Estimates page count for billing purposes.

    Billing Logic:
    - PDF: Physical page count from metadata
    - PPTX: Slide count
    - DOC/DOCX/TXT/MD/JSON: Word-based estimation (count_cn_en / 500)
    - XLS/XLSX: Row-based estimation (rows / 50)
    - Images: 1 page per image
    """

    @classmethod
    def estimate(cls, file_path: str) -> int:
        """
        Estimate page count for a file.

        Args:
            file_path: Path to the file

        Returns:
            Estimated page count (minimum 1)
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf":
                return cls._estimate_pdf(file_path)
            elif suffix == ".pptx":
                return cls._estimate_pptx(file_path)
            elif suffix == ".doc":
                return cls._estimate_doc(file_path)
            elif suffix == ".docx":
                return cls._estimate_docx(file_path)
            elif suffix == ".xls":
                return cls._estimate_xls(file_path)
            elif suffix == ".xlsx":
                return cls._estimate_xlsx(file_path)
            elif suffix in [".txt", ".md", ".json", ".fragment"]:
                return cls._estimate_text(file_path)
            elif suffix in [".png", ".jpg", ".jpeg"]:
                return 1  # Image = 1 page
            else:
                logger.warning(
                    f"Unknown file type for billing: {suffix}, defaulting to 1 page"
                )
                return 1
        except Exception as e:
            logger.error(f"Error estimating pages for {file_path}: {e}")
            return 1  # Fallback to minimum charge

    @classmethod
    def _estimate_pdf(cls, file_path: str) -> int:
        """
        Estimate pages for PDF using physical page count from metadata.
        """
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            return max(1, len(reader.pages))
        except ImportError:
            logger.warning("pypdf not installed, defaulting to 1 page")
            return 1
        except Exception as e:
            logger.error(f"PDF estimation error: {e}")
            return 1

    @classmethod
    def _estimate_pptx(cls, file_path: str) -> int:
        """
        Estimate pages for PPTX using slide count.
        """
        try:
            from pptx import Presentation

            prs = Presentation(file_path)
            return max(1, len(prs.slides))
        except ImportError:
            logger.warning("python-pptx not installed, defaulting to 1 page")
            return 1
        except Exception as e:
            logger.error(f"PPTX estimation error: {e}")
            return 1

    @classmethod
    def _estimate_docx(cls, file_path: str) -> int:
        """
        Estimate pages for DOCX using word-based counting.
        """
        try:
            from docx import Document

            doc = Document(file_path)
            total_text = ""

            # Collect text from paragraphs
            for para in doc.paragraphs:
                total_text += para.text + " "

            # Collect text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        total_text += cell.text + " "

            word_count = count_cn_en(total_text)
            return max(1, math.ceil(word_count / WORDS_PER_PAGE))

        except ImportError:
            logger.warning("python-docx not installed")
            return 1
        except Exception as e:
            logger.error(f"DOCX estimation error: {e}")
            return 1

    @classmethod
    def _estimate_doc(cls, file_path: str) -> int:
        """
        Estimate pages for DOC by converting it to DOCX first.
        """
        try:
            from app.services.document_parser.legacy_converter import doc_to_docx

            with tempfile.TemporaryDirectory(prefix="page-estimator-doc-") as temp_dir:
                converted_path, _ = doc_to_docx(file_path, temp_dir)
                return cls._estimate_docx(converted_path)
        except Exception as e:
            logger.error(f"DOC estimation error: {e}")
            return 1

    @classmethod
    def _estimate_xlsx(cls, file_path: str) -> int:
        """
        Estimate pages for XLSX using row count.
        """
        try:
            import pandas as pd

            xlsx = pd.ExcelFile(file_path)
            total_rows = 0

            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                total_rows += len(df)

            return max(1, math.ceil(total_rows / ROWS_PER_PAGE))

        except ImportError:
            logger.warning("pandas not installed")
            return 1
        except Exception as e:
            logger.error(f"XLSX estimation error: {e}")
            return 1

    @classmethod
    def _estimate_xls(cls, file_path: str) -> int:
        """
        Estimate pages for XLS by converting it to XLSX first.
        """
        try:
            from app.services.document_parser.legacy_converter import xls_to_xlsx

            with tempfile.TemporaryDirectory(prefix="page-estimator-xls-") as temp_dir:
                converted_path, _ = xls_to_xlsx(file_path, temp_dir)
                return cls._estimate_xlsx(converted_path)
        except Exception as e:
            logger.error(f"XLS estimation error: {e}")
            return 1

    @classmethod
    def _estimate_text(cls, file_path: str) -> int:
        """
        Estimate pages for text files using word-based counting.
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            word_count = count_cn_en(content)
            return max(1, math.ceil(word_count / WORDS_PER_PAGE))

        except Exception as e:
            logger.error(f"Text estimation error: {e}")
            return 1
